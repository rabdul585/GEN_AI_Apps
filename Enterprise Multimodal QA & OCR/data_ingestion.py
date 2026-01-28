"""
Data Ingestion Module
Handles document loading, chunking, embeddings, vector storage, and retrieval
"""

import json
import hashlib
import requests
import time
import re
import base64
import io
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass
from PIL import Image

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from config import *

# ============================================================================
# DOCUMENT MANAGEMENT
# ============================================================================

@dataclass
class Document:
    """Document with text and metadata"""
    text: str
    metadata: Dict
    doc_id: str = None
    
    def __post_init__(self):
        if not self.doc_id:
            content = f"{self.text}_{datetime.now().isoformat()}"
            self.doc_id = hashlib.md5(content.encode()).hexdigest()
        self.metadata["doc_id"] = self.doc_id
        self.metadata["created_at"] = datetime.now().isoformat()

# ============================================================================
# TEXT CHUNKING
# ============================================================================

class TextChunker:
    """Chunk text using various strategies"""
    
    def __init__(self, chunk_size=DEFAULT_CHUNK_SIZE, chunk_overlap=DEFAULT_CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def chunk_by_semantic(self, text: str) -> List[str]:
        """Split by paragraphs (semantic boundaries)"""
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current = []
        current_len = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_len = len(para)
            
            if current_len + para_len > self.chunk_size and current:
                chunks.append("\n\n".join(current))
                current = [para]
                current_len = para_len
            else:
                current.append(para)
                current_len += para_len
        
        if current:
            chunks.append("\n\n".join(current))
        
        return [c for c in chunks if len(c) >= MIN_CHUNK_SIZE]
    
    def chunk_by_fixed(self, text: str) -> List[str]:
        """Fixed-size chunks with overlap"""
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end].strip()
            if len(chunk) >= MIN_CHUNK_SIZE:
                chunks.append(chunk)
            start = end - self.chunk_overlap
        return chunks
    
    def chunk_by_sentence(self, text: str) -> List[str]:
        """Chunk by grouping sentences"""
        sentences = re.split(r'[.!?]+\s+', text)
        chunks = []
        current = []
        current_len = 0
        
        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue
            
            sent_len = len(sent)
            if current_len + sent_len > self.chunk_size and current:
                chunks.append(". ".join(current) + ".")
                current = [sent]
                current_len = sent_len
            else:
                current.append(sent)
                current_len += sent_len
        
        if current:
            chunks.append(". ".join(current) + ".")
        
        return [c for c in chunks if len(c) >= MIN_CHUNK_SIZE]
    
    def chunk(self, text: str, strategy: str = "Semantic") -> List[str]:
        """Chunk text using specified strategy"""
        if strategy == "Semantic":
            return self.chunk_by_semantic(text)
        elif strategy == "Fixed Size":
            return self.chunk_by_fixed(text)
        elif strategy == "Sentence-based":
            return self.chunk_by_sentence(text)
        else:
            return self.chunk_by_semantic(text)

# ============================================================================
# EMBEDDING GENERATION
# ============================================================================

class EmbeddingManager:
    """Generate embeddings for text"""
    
    def __init__(self, provider="huggingface"):
        self.provider = provider
        self.model = None
        self.dimension = 0
        
        if provider == "huggingface":
            if not SENTENCE_TRANSFORMERS_AVAILABLE:
                raise ImportError("sentence-transformers required")
            print(f"Loading embedding model...")
            self.model = SentenceTransformer(DEFAULT_EMBEDDING_MODEL)
            self.dimension = self.model.get_sentence_embedding_dimension()
            print(f"✅ Loaded {DEFAULT_EMBEDDING_MODEL} (dim: {self.dimension})")
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts"""
        if self.provider == "huggingface" and self.model:
            embeddings = self.model.encode(
                texts,
                batch_size=EMBEDDING_BATCH_SIZE,
                show_progress_bar=len(texts) > 10,
                convert_to_numpy=True
            )
            return embeddings.tolist()
        return []
    
    def embed_query(self, query: str) -> List[float]:
        """Generate embedding for single query"""
        if self.provider == "huggingface" and self.model:
            embedding = self.model.encode(query, convert_to_numpy=True)
            return embedding.tolist()
        return []

# ============================================================================
# VECTOR STORE (CHROMADB)
# ============================================================================

class VectorStore:
    """ChromaDB vector storage"""
    
    def __init__(self, embedding_manager: EmbeddingManager):
        if not CHROMADB_AVAILABLE:
            raise ImportError("chromadb required")
        
        self.embedding_manager = embedding_manager
        self.client = chromadb.PersistentClient(
            path=str(CHROMA_DB_DIR),
            settings=Settings(anonymized_telemetry=False, allow_reset=True)
        )
        
        try:
            self.collection = self.client.get_collection(CHROMA_COLLECTION_NAME)
            print(f"📚 Loaded collection: {CHROMA_COLLECTION_NAME}")
        except:
            self.collection = self.client.create_collection(
                name=CHROMA_COLLECTION_NAME,
                metadata={"hnsw:space": CHROMA_DISTANCE_METRIC}
            )
            print(f"🆕 Created collection: {CHROMA_COLLECTION_NAME}")
    
    def add_documents(self, texts: List[str], metadatas: List[Dict]) -> List[str]:
        """Add documents to vector store"""
        import uuid
        ids = [str(uuid.uuid4()) for _ in texts]
        
        print(f"Generating embeddings for {len(texts)} documents...")
        embeddings = self.embedding_manager.embed_documents(texts)
        
        # Sanitize metadata
        metadatas = [{k: str(v) if v is not None else "None" for k, v in m.items()} for m in metadatas]
        
        self.collection.add(
            ids=ids,
            documents=texts,
            embeddings=embeddings,
            metadatas=metadatas
        )
        
        print(f"✅ Added {len(texts)} documents")
        return ids
    
    def query(self, query_text: str, n_results: int = DEFAULT_RETRIEVAL_K, 
             where: Dict = None) -> Dict:
        """Query similar documents"""
        query_embedding = self.embedding_manager.embed_query(query_text)
        
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=n_results,
            where=where
        )
        
        return {
            "documents": results["documents"][0] if results["documents"] else [],
            "metadatas": results["metadatas"][0] if results["metadatas"] else [],
            "distances": results["distances"][0] if results["distances"] else []
        }
    
    def count(self) -> int:
        """Get document count"""
        return self.collection.count()
    
    def clear(self):
        """Clear all documents"""
        all_data = self.collection.get()
        if all_data["ids"]:
            self.collection.delete(ids=all_data["ids"])
        print("🗑️ Cleared collection")

# ============================================================================
# OCR PROCESSING
# ============================================================================

def image_to_base64(image_source) -> Optional[str]:
    """Convert image to base64"""
    try:
        if isinstance(image_source, (str, Path)):
            image = Image.open(image_source)
        elif hasattr(image_source, 'read'):
            image = Image.open(image_source)
        else:
            return None
        
        if image.size[0] > MAX_IMAGE_SIZE[0] or image.size[1] > MAX_IMAGE_SIZE[1]:
            image.thumbnail(MAX_IMAGE_SIZE, Image.Resampling.LANCZOS)
        
        if image.mode not in ('RGB', 'L'):
            image = image.convert('RGB')
        
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode()
    except Exception as e:
        print(f"Error: {e}")
        return None

def extract_text_from_image(image_base64: str, model_id: str = DEFAULT_OCR_MODEL) -> Dict:
    """Extract text using OpenRouter API"""
    if not OPENROUTER_API_KEY:
        return {"success": False, "error": "API key not configured"}
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": model_id,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": OCR_PROMPT},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_base64}"}}
            ]
        }]
    }
    
    for attempt in range(MAX_RETRIES):
        try:
            response = requests.post(OPENROUTER_API_URL, headers=headers, json=data, timeout=REQUEST_TIMEOUT)
            
            if response.status_code == 200:
                result = response.json()
                return {
                    "success": True,
                    "text": result['choices'][0]['message']['content']
                }
            elif response.status_code == 429:
                if attempt < MAX_RETRIES - 1:
                    wait_time = RETRY_DELAY * (2 ** attempt)
                    print(f"⏳ Rate limited. Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                    continue
                return {"success": False, "error": "Rate limit exceeded"}
            else:
                return {"success": False, "error": f"API Error: {response.status_code}"}
        except Exception as e:
            if attempt < MAX_RETRIES - 1:
                time.sleep(RETRY_DELAY)
                continue
            return {"success": False, "error": str(e)}
    
    return {"success": False, "error": "Failed after retries"}


def extract_text_from_pdf(pdf_file) -> Dict:
    """Extract text from PDF file"""
    if not PDF_SUPPORT:
        return {"success": False, "error": "PyPDF2 not installed. Run: pip install PyPDF2"}
    
    try:
        # Reset file pointer
        pdf_file.seek(0)
        
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Extract text from all pages
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")
        
        if not text_parts:
            return {"success": False, "error": "No text found in PDF"}
        
        full_text = "\n\n".join(text_parts)
        
        return {
            "success": True,
            "text": full_text,
            "pages": len(pdf_reader.pages)
        }
    
    except Exception as e:
        return {"success": False, "error": f"PDF extraction failed: {str(e)}"}


def extract_text_from_docx(docx_file) -> Dict:
    """Extract text from Word document"""
    if not DOCX_SUPPORT:
        return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}
    
    try:
        # Reset file pointer
        docx_file.seek(0)
        
        # Read DOCX
        doc = DocxDocument(docx_file)
        
        # Extract text from all paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            return {"success": False, "error": "No text found in document"}
        
        full_text = "\n\n".join(text_parts)
        
        return {
            "success": True,
            "text": full_text,
            "paragraphs": len(doc.paragraphs)
        }
    
    except Exception as e:
        return {"success": False, "error": f"DOCX extraction failed: {str(e)}"}


def extract_text_from_file(file, file_type: str = "auto") -> Dict:
    """
    Universal text extraction function
    Automatically detects file type and extracts text
    """
    # Get file extension
    filename = file.name if hasattr(file, 'name') else "unknown"
    ext = Path(filename).suffix.lower().replace('.', '')
    
    # Override with specified type
    if file_type != "auto":
        ext = file_type.lower()
    
    # Extract based on file type
    if ext in ['pdf']:
        return extract_text_from_pdf(file)
    
    elif ext in ['docx', 'doc']:
        if ext == 'doc':
            return {"success": False, "error": ".doc format not supported. Please convert to .docx"}
        return extract_text_from_docx(file)
    
    elif ext in ['png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp', 'tiff']:
        # Image - use OCR
        img_b64 = image_to_base64(file)
        if not img_b64:
            return {"success": False, "error": "Failed to convert image"}
        return extract_text_from_image(img_b64)
    
    else:
        return {"success": False, "error": f"Unsupported file type: {ext}"}

# ============================================================================
# RAG RETRIEVAL & GENERATION
# ============================================================================

class RAGPipeline:
    """Complete RAG pipeline"""
    
    def __init__(self, vector_store: VectorStore, model_id: str = DEFAULT_RAG_MODEL,
                 temperature: float = DEFAULT_TEMPERATURE, top_p: float = DEFAULT_TOP_P,
                 max_tokens: int = DEFAULT_MAX_TOKENS):
        self.vector_store = vector_store
        self.model_id = model_id
        self.temperature = temperature
        self.top_p = top_p
        self.max_tokens = max_tokens
        
        if not OPENROUTER_API_KEY:
            raise ValueError("OPENROUTER_API_KEY required")
        
        print(f"✅ RAG Pipeline initialized")
        print(f"   Model: {self.model_id}")
        print(f"   Temperature: {self.temperature}, Top-P: {self.top_p}")
    
    def retrieve_context(self, query: str, k: int = DEFAULT_RETRIEVAL_K, 
                        filter_source: str = None) -> Tuple[List[str], List[Dict]]:
        """Retrieve relevant context"""
        where = {"source_file": filter_source} if filter_source else None
        results = self.vector_store.query(query, n_results=k, where=where)
        
        context_texts = results["documents"]
        metadatas = results["metadatas"]
        
        return context_texts, metadatas
    
    def generate_response(self, query: str, context: str) -> str:
        """Generate response using LLM"""
        user_prompt = RAG_QUERY_TEMPLATE.format(context=context, question=query)
        
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": self.model_id,
            "messages": [
                {"role": "system", "content": RAG_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": self.temperature,
            "top_p": self.top_p,
            "max_tokens": self.max_tokens
        }
        
        for attempt in range(MAX_RETRIES):
            try:
                response = requests.post(OPENROUTER_API_URL, headers=headers, json=data, timeout=REQUEST_TIMEOUT)
                
                if response.status_code == 200:
                    result = response.json()
                    return result['choices'][0]['message']['content']
                elif response.status_code == 429:
                    if attempt < MAX_RETRIES - 1:
                        wait_time = RETRY_DELAY * (2 ** attempt)
                        print(f"⏳ Rate limited. Retrying in {wait_time}s...")
                        time.sleep(wait_time)
                        continue
                    raise Exception("Rate limit exceeded")
                else:
                    raise Exception(f"API Error: {response.status_code}")
            except Exception as e:
                if attempt < MAX_RETRIES - 1:
                    time.sleep(RETRY_DELAY)
                    continue
                raise
        
        raise Exception("Failed to generate response")
    
    def query(self, question: str, k: int = DEFAULT_RETRIEVAL_K, 
             filter_source: str = None) -> Dict:
        """Complete RAG query"""
        # Retrieve
        context_texts, metadatas = self.retrieve_context(question, k, filter_source)
        
        if not context_texts:
            return {
                "answer": "I couldn't find relevant information in the documents.",
                "sources": [],
                "num_sources": 0
            }
        
        # Format context
        context = "\n\n---\n\n".join([
            f"[Source: {m.get('source_file', 'Unknown')}]\n{text}"
            for text, m in zip(context_texts, metadatas)
        ])
        
        # Generate
        answer = self.generate_response(question, context)
        
        # Extract sources
        sources = []
        seen = set()
        for meta in metadatas:
            source_file = meta.get("source_file", "Unknown")
            if source_file not in seen:
                sources.append({"file": source_file})
                seen.add(source_file)
        
        return {
            "answer": answer,
            "sources": sources,
            "num_sources": len(context_texts),
            "context": context_texts
        }

# ============================================================================
# COMPLETE PIPELINE
# ============================================================================

class DataIngestionPipeline:
    """Complete end-to-end pipeline"""
    
    def __init__(self):
        self.documents = []
        self.chunker = None
        self.embedding_manager = None
        self.vector_store = None
        self.rag_pipeline = None
    
    def initialize(self, chunking_strategy="Semantic", chunk_size=DEFAULT_CHUNK_SIZE,
                  chunk_overlap=DEFAULT_CHUNK_OVERLAP, embedding_provider="huggingface",
                  rag_model=DEFAULT_RAG_MODEL, temperature=DEFAULT_TEMPERATURE,
                  top_p=DEFAULT_TOP_P, max_tokens=DEFAULT_MAX_TOKENS, retrieval_k=DEFAULT_RETRIEVAL_K):
        """Initialize complete pipeline"""
        print("\n🚀 Initializing Data Ingestion Pipeline...")
        
        self.chunker = TextChunker(chunk_size, chunk_overlap)
        print(f"✅ Chunker: {chunking_strategy}")
        
        self.embedding_manager = EmbeddingManager(embedding_provider)
        print(f"✅ Embeddings: {embedding_provider}")
        
        self.vector_store = VectorStore(self.embedding_manager)
        print(f"✅ Vector Store: ChromaDB")
        
        self.rag_pipeline = RAGPipeline(
            self.vector_store,
            model_id=rag_model,
            temperature=temperature,
            top_p=top_p,
            max_tokens=max_tokens
        )
        print(f"✅ RAG Pipeline ready\n")
    
    def process_document(self, file, ocr_model: str, chunking_strategy: str) -> Dict:
        """Process document: Extract Text → Chunk → Embed → Store"""
        # Get file extension
        filename = file.name if hasattr(file, 'name') else "unknown"
        ext = Path(filename).suffix.lower().replace('.', '')
        
        # Extract text based on file type
        if ext in ['pdf']:
            result = extract_text_from_pdf(file)
        elif ext in ['docx']:
            result = extract_text_from_docx(file)
        elif ext in ['png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp', 'tiff']:
            # Image - use OCR
            img_b64 = image_to_base64(file)
            if not img_b64:
                return {"success": False, "error": "Image conversion failed"}
            result = extract_text_from_image(img_b64, ocr_model)
        else:
            return {"success": False, "error": f"Unsupported file type: {ext}"}
        
        if not result["success"]:
            return result
        
        text = result["text"]
        
        if len(text.strip()) < MIN_TEXT_LENGTH:
            return {"success": False, "error": "Text too short"}
        
        # Create document
        doc = Document(
            text=text,
            metadata={
                "source_file": filename,
                "file_type": ext,
                "extraction_method": "pdf" if ext == "pdf" else "docx" if ext == "docx" else "ocr",
                "ocr_model": ocr_model if ext in ['png', 'jpg', 'jpeg', 'gif', 'webp'] else "N/A",
                "text_length": len(text),
                "timestamp": datetime.now().isoformat(),
                "pages": result.get("pages", 1)
            }
        )
        self.documents.append(doc)
        
        # Chunk
        chunks = self.chunker.chunk(text, chunking_strategy)
        
        # Add metadata to chunks
        chunk_metadatas = [
            {**doc.metadata, "chunk_index": i}
            for i in range(len(chunks))
        ]
        
        # Store in vector DB
        self.vector_store.add_documents(chunks, chunk_metadatas)
        
        return {
            "success": True,
            "doc_id": doc.doc_id,
            "chunks": len(chunks),
            "text_length": len(text),
            "file_type": ext
        }
    
    def query_documents(self, question: str, k: int = None, filter_source: str = None) -> Dict:
        """Query documents"""
        if not self.rag_pipeline:
            return {"error": "Pipeline not initialized"}
        
        k = k or DEFAULT_RETRIEVAL_K
        return self.rag_pipeline.query(question, k, filter_source)

# ============================================================================
# UTILITIES
# ============================================================================

def format_timestamp() -> str:
    """Format current timestamp"""
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def truncate_text(text: str, max_length: int = 100) -> str:
    """Truncate text"""
    return text[:max_length] + "..." if len(text) > max_length else text